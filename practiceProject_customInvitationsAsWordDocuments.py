# Program that generates a Word document with custom invitations based on the guests.txt file.

import docx

guestList = [] # list of guests on guests.txt file.

# Read .txt file and extract the guest names:
guests = open('<insert path here>') # This is the path to where the file guests.txt should be.
for line in guests:
    guestList.append(line) # each line is stored in the list guestList.
guests.close()
# Use the guest name and create one invitation per Word page:
print('Creating an invitation document...')
doc = docx.Document()
for i in guestList:
    doc.add_paragraph("It would be a pleasure to have the company of \n" + i + "at 11010 Memory Lane on the Evening of \n" + "April 1st \n" + "at 7 o' clock")
    for x in range(len(doc.paragraphs)):
        doc.paragraphs[x].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE) # adds a page break after writing the invitation for each guest.
# Save the document with the invitations:
doc.save('Invitations.docx')
print('Done.')
